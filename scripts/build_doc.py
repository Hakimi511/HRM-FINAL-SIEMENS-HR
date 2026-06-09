# -*- coding: utf-8 -*-
"""生成项目说明 Word 文档"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SHOTS = os.path.join(ROOT, 'shots')

PETROL = RGBColor(0x00, 0x76, 0x80)
PETROL_BRIGHT = RGBColor(0x00, 0x99, 0x99)
DEEP = RGBColor(0x00, 0x10, 0x2A)
GREY = RGBColor(0x6B, 0x72, 0x80)
ORANGE = RGBColor(0xEC, 0x66, 0x02)
CJK = 'Microsoft YaHei'

doc = Document()

# ---------- 全局样式 ----------
def style_font(style, name=CJK, size=None, bold=None, color=None):
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn('w:eastAsia'), name)
    rf.set(qn('w:ascii'), name)
    rf.set(qn('w:hAnsi'), name)
    if size is not None: style.font.size = Pt(size)
    if bold is not None: style.font.bold = bold
    if color is not None: style.font.color.rgb = color

style_font(doc.styles['Normal'], size=10.5)
style_font(doc.styles['Heading 1'], size=17, bold=True, color=PETROL)
style_font(doc.styles['Heading 2'], size=13.5, bold=True, color=DEEP)
style_font(doc.styles['Heading 3'], size=11.5, bold=True, color=PETROL)
for s in ['Title', 'Subtitle']:
    try: style_font(doc.styles[s])
    except KeyError: pass

# 页边距
for sec in doc.sections:
    sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.4); sec.right_margin = Cm(2.4)

IMG_W = Inches(6.3)

# ---------- 辅助 ----------
def run_cjk(p, text, size=10.5, bold=False, color=None, italic=False):
    r = p.add_run(text)
    r.font.name = CJK; r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    rpr = r._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    rf.set(qn('w:eastAsia'), CJK)
    if color is not None: r.font.color.rgb = color
    return r

def para(text='', size=10.5, bold=False, color=None, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.25
    if text: run_cjk(p, text, size=size, bold=bold, color=color)
    return p

def h1(text):
    p = doc.add_heading(level=1); run_cjk(p, text, size=17, bold=True, color=PETROL)
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(8)
    return p

def h2(text):
    p = doc.add_heading(level=2); run_cjk(p, text, size=13.5, bold=True, color=DEEP)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(5)
    return p

def h3(text):
    p = doc.add_heading(level=3); run_cjk(p, text, size=11.5, bold=True, color=PETROL)
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(3)
    return p

def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.22
    if bold_lead:
        run_cjk(p, bold_lead, bold=True)
        run_cjk(p, text)
    else:
        run_cjk(p, text)
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.22
    run_cjk(p, text)
    return p

def add_image(name, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(os.path.join(SHOTS, name), width=IMG_W)
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)
    run_cjk(c, caption, size=9, italic=True, color=GREY)

def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def fill_cell(cell, text, bold=False, color=None, size=10, white=False):
    cell.text = ''
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    run_cjk(p, text, size=size, bold=bold, color=(RGBColor(0xFF,0xFF,0xFF) if white else color))

def make_table(headers, rows, widths, header_bg='00767F'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, hd in enumerate(headers):
        fill_cell(t.rows[0].cells[i], hd, bold=True, white=True, size=10)
        set_cell_bg(t.rows[0].cells[i], header_bg)
        t.rows[0].cells[i].width = widths[i]
    for r, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            fill_cell(cells[i], val, size=9.5)
            cells[i].width = widths[i]
            if r % 2 == 1: set_cell_bg(cells[i], 'F2F6F7')
    # 设置列宽（需逐单元格）
    for row in t.rows:
        for i, c in enumerate(row.cells):
            c.width = widths[i]
    return t

# =====================================================================
# 封面
# =====================================================================
band = para('SIEMENS  ·  SMART INFRASTRUCTURE', size=11, bold=True, color=PETROL_BRIGHT, space_after=2)
para('渠道销售（SI）培训考核', size=9, color=GREY, space_after=18)

tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_cjk(tp, 'SI 渠道销售培训考核学习平台', size=26, bold=True, color=DEEP)
tp.paragraph_format.space_after = Pt(4)
sp = doc.add_paragraph()
run_cjk(sp, 'AI 赋能 · “学习—考核—分析”一体化 Web 平台', size=13, bold=True, color=PETROL)
sp.paragraph_format.space_after = Pt(2)
para('项目说明文档', size=12, color=GREY, space_after=16)

# 信息表
info = make_table(
    ['项目信息', '内容'],
    [
        ['项目名称', '西门子 SI 渠道销售培训考核学习平台'],
        ['适用对象', '西门子 SI 渠道销售伙伴（学员）｜HR / 培训管理员'],
        ['项目网址', 'https://hakimi511.github.io/HRM-FINAL-SIEMENS-HR/'],
        ['GitHub 仓库', 'https://github.com/Hakimi511/HRM-FINAL-SIEMENS-HR'],
        ['HR 后台演示口令', 'siemens2025'],
        ['提交部门 / 日期', 'HR 部门　｜　2026 年'],
    ],
    [Cm(4.2), Cm(11.2)],
)
para('', space_after=2)
para('提示：本平台为纯前端原型，所有功能可在上述网址直接体验；学员端无需登录，HR 后台以演示口令进入。',
     size=9, color=GREY)

doc.add_page_break()

# =====================================================================
# 一、项目背景
# =====================================================================
h1('一、项目背景：主要解决的问题')

h2('1.1 业务场景')
para('西门子智能基础设施集团（Smart Infrastructure，SI）的大量产品——包括 EA 中压产品线'
     '（气体/空气绝缘开关设备、保护测控装置、配电数字化平台）与 EP 低压产品线（断路器、'
     '成套开关设备、电柜数字化）——主要通过 SI 渠道销售伙伴（系统集成商、经销商）触达终端客户。'
     '这些渠道销售人员数量多、地域分散、流动性强，且面对的产品技术复杂、数字化解决方案更新快，'
     '既要懂产品，又要会销售、谈判与客户管理。')

h2('1.2 现存痛点')
bullet('培训资料分散，缺乏体系。', '学习内容散落：')
bullet('考核手段薄弱，缺乏标准化、可量化的知识测评，学得好不好全凭感觉。', '以学代考：')
bullet('学习效果不可见：HR / 培训经理无法掌握“谁学了、学得怎样、哪里薄弱”。', '过程黑箱：')
bullet('培训决策缺数据支撑，不清楚该重点补哪条产品线 / 哪个知识点、哪个区域团队最弱。', '决策无据：')
bullet('传统 LMS 系统重、上线慢、成本高，难以快速落地一个轻量、好用的工具。', '工具笨重：')

h2('1.3 项目目标')
para('用一个轻量、零部署成本、即开即用的 Web 平台，把“学（资料）— 考（题库）— 析（数据）”打通成闭环：'
     '对销售学员，提供随时随地的自学与自测；对 HR，提供一个掌握全员学习考核情况的数据驾驶舱，'
     '让培训从“凭经验”走向“看数据”。')

# =====================================================================
# 二、解决思路
# =====================================================================
h1('二、解决思路：项目设计与逻辑')

h2('2.1 总体理念：以考促学 + 数据驱动')
para('平台围绕三大闭环展开，形成从个人学习到组织决策的完整链路：')
numbered('学习闭环：浏览产品资料 → 在线考核检验 → 错题精练巩固。')
numbered('数据闭环：每次答题记录“逐题明细” → 沉淀为个人/组织数据 → 支撑多维分析。')
numbered('管理闭环：HR 查看分析 → 定位薄弱环节与区域差异 → 精准投放培训资源。')

h2('2.2 双角色 · 双模式')
make_table(
    ['模式', '进入方式', '核心能力'],
    [
        ['学员模式（默认）', '首次访问登记姓名 + 销售区域', '在线考核、学习资料、个人成绩与错题本、成绩导出'],
        ['HR 管理后台', '顶部入口 + 口令 siemens2025', '全员学习考核数据分析、薄弱点定位、报表导入导出'],
    ],
    [Cm(3.6), Cm(5.2), Cm(6.6)],
)

h2('2.3 内容架构')
para('题库与资料均由真实业务内容整理而来，并做结构化归并：', space_after=4)
make_table(
    ['维度', '规模', '说明'],
    [
        ['考核题库', '136 题 / 25 知识分类', '单选 87、多选 49（含判断题），归并为 5 大能力模块'],
        ['能力模块', '5 个', '数字化产品、销售技能、谈判与沟通、职场软技能、财务知识'],
        ['学习资料', '20 份 PDF / 12 分类', 'EA（中压）与 EP（低压）两大产品线销售资料'],
    ],
    [Cm(3.4), Cm(4.6), Cm(7.4)],
)

h2('2.4 系统设计与关键技术决策')
bullet('原生 HTML / CSS / JavaScript（ES Modules），零构建、零依赖，托管于 GitHub Pages，零成本、秒级部署、易交付维护。', '纯前端静态站：')
bullet('题库 Excel 自动转 questions.json、PDF 目录自动转 materials.json（由 build_data.py 生成），内容更新只需替换源文件并重跑脚本。', '数据驱动：')
bullet('成绩、错题、学习进度保存在浏览器本地（localStorage），无需后端、保护隐私；跨设备数据通过“学员导出 → HR 导入”在 HR 端汇总。', '本地存储 + 导入导出：')
bullet('每次答题记录每道题的对错与所属模块/分类，为部门对比、知识薄弱点、难题排行等分析提供数据基础。', '逐题明细：')

h2('2.5 数据流')
para('学员答题 → 生成成绩 + 逐题明细 → 写入本地存储 → 一键导出 JSON → HR 导入汇总 → '
     '按学员 / 部门 / 模块 / 分类 / 题目多维聚合 → 输出 KPI、图表与培训建议 → 支撑培训决策。',
     space_after=4)

h2('2.6 考核引擎逻辑')
bullet('组卷：范围（综合/模块/分类/错题本）× 题量 × 模式（考试/练习）× 是否计时。')
bullet('答题：单选/多选/判断自适应，选项可乱序并自动重映射答案，配答题卡与计时。')
bullet('判分：合格线 60%，自动评级（优秀 ≥90 / 良好 ≥80 / 合格 ≥60 / 不合格）。')
bullet('巩固：答错自动入错题本，答对自动移出，支持错题专项重练与逐题解析回顾。')

h2('2.7 分析逻辑')
bullet('多维聚合：按学员、部门、能力模块、知识分类、单题分别统计正确率。')
bullet('薄弱定位：正确率低于 70% 自动标红，并据此生成“优先培训方向”建议。')
bullet('难题洞察：统计全员错得最多的题目，反推内容难点与表述问题。')

# =====================================================================
# 三、功能演示
# =====================================================================
doc.add_page_break()
h1('三、项目功能演示')
para('以下为平台真实运行界面截图（演示数据），完整体验请访问项目网址。', size=9.5, color=GREY)

h2('3.1 学员端 · 首页')
para('统一的西门子品牌视觉，呈现题库与资料规模概览、考核/学习两大入口、五大能力模块，以及个人学习概览。')
add_image('home.png', '图 1　学员端首页：概览、入口与五大能力模块')

h2('3.2 考核中心')
para('支持按综合题库 / 能力模块 / 知识分类 / 错题本灵活组卷，自选题量、考试或练习模式、是否计时；'
     '答题界面配有答题卡、进度与计时，交卷后自动判分并提供逐题解析。')
add_image('exam.png', '图 2　考核中心：组卷设置（范围 / 题量 / 模式 / 计时）')

h2('3.3 学习资料库')
para('EA / EP 两大产品线、12 个产品分类、20 份销售资料，支持在线预览、下载、型号关键词搜索与“已读”标记。')
add_image('library.png', '图 3　学习资料库：EA / EP 产品资料在线预览与下载')

h2('3.4 我的成绩（学员）')
para('个人 KPI（考核次数、平均/最高正确率、合格次数）、历史记录、错题本，并支持一键导出成绩交 HR 汇总。')
add_image('records.png', '图 4　我的成绩：个人学习档案与成绩导出')

h2('3.5 HR 管理后台')
para('口令进入。汇集全员数据，提供达标概览、各销售区域对比、知识薄弱点分析与培训建议、难题排行、'
     '学员排行榜与明细表，并支持成绩导入、CSV / JSON 报表导出与一键演示数据。')
add_image('admin_top.png', '图 5　HR 管理后台（上）：KPI 总览、达标环形、区域对比与知识薄弱点分析')
add_image('admin_bottom.png', '图 6　HR 管理后台（下）：薄弱分类、难题排行、排行榜与学员明细表')

h2('3.6 体验入口与演示动线')
bullet('项目网址：https://hakimi511.github.io/HRM-FINAL-SIEMENS-HR/')
bullet('演示动线：首页登记身份 → 进入“HR 管理后台”输入口令 siemens2025 → 点“载入演示数据”，'
       '即可看到完整的数据分析驾驶舱。')

# =====================================================================
# 四、价值与风险
# =====================================================================
doc.add_page_break()
h1('四、可能的价值与风险')

h2('4.1 项目价值')
h3('对销售学员')
bullet('随时随地自学自测、查漏补缺、错题精练，系统提升产品知识与销售赢单能力。')
h3('对 HR / 培训管理')
bullet('量化学习效果，直观定位薄弱知识点与区域团队差异，数据驱动培训资源精准投放，显著降低培训管理成本。')
h3('对组织')
bullet('沉淀标准化知识与培训资产，加速新人上手，提升渠道整体专业度，最终促进销售转化。')
h3('技术与 AI 价值')
bullet('零成本、秒级部署、易维护，可快速复制到其他产品线或业务场景。')
bullet('本项目从数据清洗、内容结构化、前端开发到部署上线，全流程由 AI 辅助完成，是 AI 提效的实践范例。')

h2('4.2 风险与应对')
make_table(
    ['潜在风险', '说明', '应对建议'],
    [
        ['数据未中心化', '纯前端无后端，跨设备/规模化统计依赖导入导出', '正式落地对接企业 LMS / 数据库或 Serverless + 云数据库实现集中统计'],
        ['身份与权限为前端模拟', '口令为前端校验，非真实鉴权，仅适合原型/演示', '接入企业 AD / SSO 与后端权限管理'],
        ['数据安全与合规', 'localStorage 明文存储，涉及员工个人信息', '后端化 + 传输加密 + 权限分级，遵循个人信息保护要求'],
        ['内容时效性', '产品资料与题库需持续更新', '建立内容运营机制，脚本化批量更新数据'],
        ['题库质量', '部分题目暂缺解析、难度维度单一', '补充解析、分级难度、扩充题量'],
        ['浏览器缓存', '更新后老用户可能看到旧版本', '发布加版本号 / 配置缓存策略，提示强制刷新'],
        ['内容版权', '题库与资料为西门子内部资产', '限内部访问，配合访问控制与水印'],
    ],
    [Cm(3.2), Cm(6.0), Cm(6.2)],
)

h2('4.3 后续规划')
bullet('后端化：引入轻量后端 + 数据库，实现集中统计、SSO 登录与权限管理。')
bullet('AI 增强：基于知识库的 AI 问答助手、智能出题与个性化学习路径推荐。')
bullet('内容扩展：覆盖更多产品线与销售场景，建立常态化题库/资料更新机制。')
bullet('激励体系：引入学习积分、勋章与排行榜，提升学员参与度。')

para('', space_after=4)
end = para('— 本文档由 AI 辅助生成，配合项目网址与 GitHub 仓库共同提交 —', size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

out = os.path.join(ROOT, '西门子SI渠道销售培训考核平台_项目说明.docx')
doc.save(out)
print('SAVED:', out)
